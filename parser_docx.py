"""
docx_parser.py — A comprehensive Python DOCX parser toolkit.

Covers:
  - Text extraction (plain, structured, with styles)
  - Table extraction
  - Metadata & document properties
  - Headings / outline
  - Images extraction
  - Comments & tracked changes
  - Hyperlinks
  - Export to Markdown / JSON

Dependencies:
    pip install python-docx lxml Pillow
"""

from __future__ import annotations

import json
import re
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


# ─────────────────────────────────────────────
# Data models
# ─────────────────────────────────────────────

@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    font_name: str | None = None
    font_size: float | None = None  # in points
    color: str | None = None        # hex RGB e.g. "FF0000"


@dataclass
class Para:
    runs: list[Run] = field(default_factory=list)
    style: str = "Normal"
    alignment: str | None = None
    level: int | None = None        # heading level if applicable

    @property
    def text(self) -> str:
        return "".join(r.text for r in self.runs)


@dataclass
class Cell:
    paragraphs: list[Para] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs)


@dataclass
class DocTable:
    rows: list[list[Cell]] = field(default_factory=list)

    def to_plain(self) -> list[list[str]]:
        return [[cell.text for cell in row] for row in self.rows]


@dataclass
class Hyperlink:
    text: str
    url: str
    paragraph_index: int


@dataclass
class Comment:
    id: str
    author: str
    date: str
    text: str


@dataclass
class ParsedDoc:
    paragraphs: list[Para] = field(default_factory=list)
    tables: list[DocTable] = field(default_factory=list)
    hyperlinks: list[Hyperlink] = field(default_factory=list)
    comments: list[Comment] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    images: list[Path] = field(default_factory=list)


# ─────────────────────────────────────────────
# Core parser
# ─────────────────────────────────────────────

class DocxParser:
    """
    Full-featured .docx parser.

    Usage
    -----
    parser = DocxParser("report.docx")
    doc    = parser.parse()

    # Plain text
    print(parser.plain_text())

    # All tables as list-of-lists
    for table in parser.tables_as_plain():
        print(table)

    # Headings / outline
    for h in parser.headings():
        print(h)

    # Export
    parser.to_markdown("report.md")
    parser.to_json("report.json")
    """

    def __init__(self, path: str | Path):
        self.path = Path(path)
        if not self.path.exists():
            raise FileNotFoundError(self.path)
        self._doc = Document(str(self.path))
        self._parsed: ParsedDoc | None = None

    # ── Public high-level API ──────────────────

    def parse(self) -> ParsedDoc:
        """Parse everything and return a ParsedDoc dataclass."""
        if self._parsed is not None:
            return self._parsed

        self._parsed = ParsedDoc(
            paragraphs=self._parse_paragraphs(),
            tables=self._parse_tables(),
            hyperlinks=self._parse_hyperlinks(),
            comments=self._parse_comments(),
            metadata=self._parse_metadata(),
        )
        return self._parsed

    def plain_text(self, separator: str = "\n") -> str:
        """Return all paragraph text joined by *separator*."""
        return separator.join(
            p.text for p in self.parse().paragraphs if p.text.strip()
        )

    def headings(self) -> list[dict]:
        """Return a list of {level, text} dicts for all heading paragraphs."""
        result = []
        for p in self.parse().paragraphs:
            if p.level is not None:
                result.append({"level": p.level, "text": p.text})
        return result

    def tables_as_plain(self) -> list[list[list[str]]]:
        """Return all tables as nested plain-text lists."""
        return [t.to_plain() for t in self.parse().tables]

    def extract_images(self, output_dir: str | Path = ".") -> list[Path]:
        """
        Extract all embedded images to *output_dir*.
        Returns list of saved Paths.
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        saved: list[Path] = []

        with zipfile.ZipFile(self.path) as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    filename = Path(name).name
                    dest = output_dir / filename
                    dest.write_bytes(z.read(name))
                    saved.append(dest)

        self.parse().images = saved
        return saved

    def to_markdown(self, output_path: str | Path | None = None) -> str:
        """Convert document to Markdown. Optionally write to *output_path*."""
        doc = self.parse()
        lines: list[str] = []

        # Paragraphs (interleaved with tables based on original order)
        body = self._doc.element.body
        table_index = 0
        para_index = 0

        for child in body:
            tag = etree.QName(child).localname
            if tag == "p":
                if para_index < len(doc.paragraphs):
                    p = doc.paragraphs[para_index]
                    lines.append(self._para_to_md(p))
                    para_index += 1
            elif tag == "tbl":
                if table_index < len(doc.tables):
                    lines.append(self._table_to_md(doc.tables[table_index]))
                    table_index += 1

        md = "\n".join(lines).strip()
        if output_path:
            Path(output_path).write_text(md, encoding="utf-8")
        return md

    def to_json(self, output_path: str | Path | None = None) -> str:
        """Serialize the parsed document to JSON."""
        doc = self.parse()

        def serialise(obj):
            if hasattr(obj, "__dataclass_fields__"):
                return asdict(obj)
            if isinstance(obj, Path):
                return str(obj)
            raise TypeError(type(obj))

        data = asdict(doc)
        # Convert Path objects
        data["images"] = [str(p) for p in doc.images]
        js = json.dumps(data, indent=2, default=str)
        if output_path:
            Path(output_path).write_text(js, encoding="utf-8")
        return js

    # ── Internal parsers ───────────────────────

    def _parse_paragraphs(self) -> list[Para]:
        return [self._parse_paragraph(p) for p in self._doc.paragraphs]

    def _parse_paragraph(self, p: Paragraph) -> Para:
        style_name: str = p.style.name if p.style else "Normal"

        # Detect heading level
        level: int | None = None
        m = re.match(r"Heading (\d+)", style_name, re.IGNORECASE)
        if m:
            level = int(m.group(1))

        runs = [self._parse_run(r) for r in p.runs]

        alignment = None
        if p.alignment is not None:
            alignment = str(p.alignment)

        return Para(runs=runs, style=style_name, alignment=alignment, level=level)

    @staticmethod
    def _parse_run(r) -> Run:
        font = r.font
        color = None
        try:
            if font.color and font.color.rgb:
                color = str(font.color.rgb)
        except Exception:
            pass

        size = None
        try:
            if font.size:
                size = font.size.pt
        except Exception:
            pass

        return Run(
            text=r.text,
            bold=bool(r.bold),
            italic=bool(r.italic),
            underline=bool(r.underline),
            font_name=font.name,
            font_size=size,
            color=color,
        )

    def _parse_tables(self) -> list[DocTable]:
        tables: list[DocTable] = []
        for tbl in self._doc.tables:
            rows: list[list[Cell]] = []
            for row in tbl.rows:
                cells: list[Cell] = []
                for cell in row.cells:
                    paras = [self._parse_paragraph(p) for p in cell.paragraphs]
                    cells.append(Cell(paragraphs=paras))
                rows.append(cells)
            tables.append(DocTable(rows=rows))
        return tables

    def _parse_hyperlinks(self) -> list[Hyperlink]:
        links: list[Hyperlink] = []
        rels = self._doc.part.rels

        for para_idx, para in enumerate(self._doc.paragraphs):
            for hl in para._p.findall(qn("w:hyperlink")):
                r_id = hl.get(qn("r:id"))
                url = rels[r_id].target_ref if r_id in rels else "#"
                text = "".join(
                    t.text for t in hl.iter(qn("w:t")) if t.text
                )
                links.append(Hyperlink(text=text, url=url, paragraph_index=para_idx))

        return links

    def _parse_comments(self) -> list[Comment]:
        comments: list[Comment] = []
        try:
            comments_part = self._doc.part.package.part_related_by(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
            )
            root = comments_part._element
        except Exception:
            # No comments part in the document
            return comments

        for c in root.findall(qn("w:comment")):
            text = "".join(
                t.text for t in c.iter(qn("w:t")) if t.text
            )
            comments.append(Comment(
                id=c.get(qn("w:id"), ""),
                author=c.get(qn("w:author"), ""),
                date=c.get(qn("w:date"), ""),
                text=text,
            ))

        return comments

    def _parse_metadata(self) -> dict[str, Any]:
        cp = self._doc.core_properties

        def _get(attr):
            val = getattr(cp, attr, None)
            return str(val) if val is not None else None

        return {
            "title":            _get("title"),
            "author":           _get("author"),
            "created":          _get("created"),
            "modified":         _get("modified"),
            "last_modified_by": _get("last_modified_by"),
            "subject":          _get("subject"),
            "description":      _get("comments"),   # 'description' in OOXML = 'comments' in python-docx
            "keywords":         _get("keywords"),
            "revision":         _get("revision"),
            "category":         _get("category"),
        }

    # ── Markdown helpers ───────────────────────

    @staticmethod
    def _para_to_md(p: Para) -> str:
        if p.level is not None:
            return f"{'#' * p.level} {p.text}"

        # Style-based mapping
        lower = p.style.lower()
        if "quote" in lower:
            return f"> {p.text}"
        if "code" in lower or "verbatim" in lower:
            return f"`{p.text}`"

        # Inline formatting (run-level)
        parts: list[str] = []
        for r in p.runs:
            t = r.text
            if not t:
                continue
            if r.bold and r.italic:
                t = f"***{t}***"
            elif r.bold:
                t = f"**{t}**"
            elif r.italic:
                t = f"*{t}*"
            if r.underline:
                t = f"<u>{t}</u>"
            parts.append(t)

        return "".join(parts)

    @staticmethod
    def _table_to_md(table: DocTable) -> str:
        plain = table.to_plain()
        if not plain:
            return ""

        def _escape(s: str) -> str:
            return s.replace("|", "\\|").replace("\n", " ")

        header = "| " + " | ".join(_escape(c) for c in plain[0]) + " |"
        separator = "| " + " | ".join("---" for _ in plain[0]) + " |"
        body = [
            "| " + " | ".join(_escape(c) for c in row) + " |"
            for row in plain[1:]
        ]
        return "\n".join([header, separator] + body)


# ─────────────────────────────────────────────
# Convenience functions
# ─────────────────────────────────────────────

def extract_text(path: str | Path) -> str:
    """Quickest way to get all text from a .docx file."""
    return DocxParser(path).plain_text()


def extract_tables(path: str | Path) -> list[list[list[str]]]:
    """Return all tables from a .docx as plain text nested lists."""
    return DocxParser(path).tables_as_plain()


def extract_headings(path: str | Path) -> list[dict]:
    """Return document outline as list of {level, text} dicts."""
    return DocxParser(path).headings()


def docx_to_markdown(path: str | Path, output: str | Path | None = None) -> str:
    """Convert a .docx to Markdown string (and optionally save it)."""
    return DocxParser(path).to_markdown(output)


def docx_to_json(path: str | Path, output: str | Path | None = None) -> str:
    """Serialize a .docx to a JSON string (and optionally save it)."""
    return DocxParser(path).to_json(output)


# ─────────────────────────────────────────────
# CLI entry point
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="Parse a .docx file")
    parser.add_argument("file", help="Path to .docx file")
    parser.add_argument(
        "--mode",
        choices=["text", "headings", "tables", "metadata", "markdown", "json"],
        default="text",
    )
    parser.add_argument("--output", "-o", help="Output file path (optional)")
    args = parser.parse_args()

    p = DocxParser(args.file)

    if args.mode == "text":
        result = p.plain_text()
    elif args.mode == "headings":
        result = json.dumps(p.headings(), indent=2)
    elif args.mode == "tables":
        result = json.dumps(p.tables_as_plain(), indent=2)
    elif args.mode == "metadata":
        result = json.dumps(p.parse().metadata, indent=2, default=str)
    elif args.mode == "markdown":
        result = p.to_markdown(args.output)
    elif args.mode == "json":
        result = p.to_json(args.output)
    else:
        result = ""

    if args.output and args.mode not in ("markdown", "json"):
        Path(args.output).write_text(result, encoding="utf-8")
    else:
        print(result)